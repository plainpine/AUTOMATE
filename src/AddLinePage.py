import docx
doc = docx.Document()
doc.add_paragraph('これは１ページ目！')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('これは２ページ目！')
doc.save('twoPage.docx')
