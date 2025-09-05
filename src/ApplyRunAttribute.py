import docx
doc = docx.Document('demo.docx')
doc.paragraphs[0].text
print(doc.paragraphs[0].text)

doc.paragraphs[0].style  # The exact id may be different:
print(doc.paragraphs[0].style)

doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text
print(doc.paragraphs[1].text)

(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)

doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')
