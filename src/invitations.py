#! python3
# invitations.py
import docx
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

doc = docx.Document()
names = open('guests.txt',encoding='utf-8')
names = names.read()
names = names.split(sep='\n')
for name in names:
    doc.add_paragraph(name).style = 'Body Text'
    doc.add_paragraph('拝啓　時下ますますご盛栄のこととお喜び申し上げます。').style = 'Body Text'
    doc.add_paragraph('このたび下記の通りパーティを開催いたしますので、ご出席賜りますようよろしくお願い申し上げます。').style = 'Body Text'
    doc.add_paragraph('敬具').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('記').alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph(' 日時：４月１日　19:00～')
    paragraph.style = 'Body Text'
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph = doc.add_paragraph(' 場所：ホテル・オライリー')
    paragraph.style = 'Body Text'
    paragraph.paragraph_format.left_indent = Inches(0.25)
    doc.add_paragraph('以上').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

doc.save('invitations.docx')
