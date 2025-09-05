import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')

para_obj_1 = doc.add_paragraph('これは第２段落です。')
para_obj_2 = doc.add_paragraph('これは第３段落です。')
para_obj_1.add_run(' これは第２段落に追加したテキストです。')

doc.save('multipleParagraphs.docx')
