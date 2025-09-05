import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')

para3 = doc.add_paragraph('これは第３段落です。')
para3.insert_paragraph_before('これは第２段落です。')
doc.add_paragraph('これは第４段落です。')

doc.save('insertParagraphs.docx')
